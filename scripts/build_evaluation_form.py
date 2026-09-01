#!/usr/bin/env python3
"""Build a reviewer-ready, editable and printable evaluation form."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "build"
BUILD.mkdir(exist_ok=True)
OUTPUT = BUILD / "ESG-DP-2026-evaluation-form.docx"
DOC_ID = "ESG-DP-2026-REVIEW"

CRITERIA = [
    (
        "1. Vymezení problému, cíl a výzkumné otázky",
        "Posuďte jasnost problému, vhodnost cíle, provázání hlavní a dílčích otázek a přiměřenost rozsahu práce.",
    ),
    (
        "2. Teoretická východiska a práce s literaturou",
        "Posuďte výběr zdrojů, porozumění teoriím, schopnost syntézy, práci s protiargumenty a přesnost citací.",
    ),
    (
        "3. Regulatorní a věcný kontext",
        "Posuďte správnost a přiměřenost výkladu CSRD, ESRS, dvojí materiality, Taxonomie a assurance. Rukopis se nevydává za právní stanovisko.",
    ),
    (
        "4. Metodologie",
        "Posuďte vhodnost dokumentové a obsahové analýzy, výběr případů, vymezení jednotek, kódovací rámec, kontrolní postupy, reflexivitu a limity.",
    ),
    (
        "5. Empirický korpus a auditovatelnost",
        "Posuďte dostatečnost dokumentů, přesnost lokátorů, dohledatelnost tvrzení, oddělení firemního sdělení od interpretace a přiměřenost škály E0-E4.",
    ),
    (
        "6. Výsledky",
        "Posuďte, zda výsledky vycházejí z dat, obsahují mezipřípadové kontrasty a negativní důkazy a nepřekračují možnosti dokumentového designu.",
    ),
    (
        "7. Diskuse a přínos",
        "Posuďte návrat k teoriím, kvalitu interpretace, originalitu přínosu, praktická doporučení a přiznání alternativních výkladů.",
    ),
    (
        "8. Jazyk, struktura a formální úroveň",
        "Posuďte srozumitelnost, akademický styl, soudržnost, míru opakování, tabulky, citace a celkovou čitelnost.",
    ),
]


def set_a4(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_table_borders(table, value: str = "single", size: str = "6", color: str = "B7B7B7") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), value)
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def setup(doc: Document) -> None:
    set_a4(doc.sections[0])
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    doc.core_properties.title = "Formulář nezávislého posudku"
    doc.core_properties.subject = f"Zaslepené hodnocení rukopisu {DOC_ID}"
    doc.core_properties.author = "Anonymizováno pro účely experimentu"
    doc.core_properties.last_modified_by = "Anonymizováno"

    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_ID} | formulář posudku")
    run.font.name = "Arial"
    run.font.size = Pt(8)

    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def field_table(doc: Document) -> None:
    items = [
        ("Identifikátor rukopisu", DOC_ID),
        ("Hodnotitel", ""),
        ("Instituce / odborná oblast", ""),
        ("Datum převzetí", ""),
        ("Datum uzavření posudku", ""),
        ("Přibližný čas věnovaný hodnocení", ""),
    ]
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for row, (label, value) in zip(table.rows, items):
        prevent_row_split(row)
        row.cells[0].width = Cm(6.0)
        row.cells[1].width = Cm(11.0)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "EDEDED")
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        r0.bold = True
        p1 = row.cells[1].paragraphs[0]
        p1.add_run(value if value else " ")
        row.height = Cm(0.8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="7F7F7F")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    prevent_row_split(table.rows[0])


def rating_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color="A6A6A6")
    labels = ["Hodnocení", "1", "2", "3", "4", "5"]
    checks = ["Zakroužkujte / označte", "[ ]", "[ ]", "[ ]", "[ ]", "[ ]"]
    for i, values in enumerate((labels, checks)):
        row = table.rows[i]
        prevent_row_split(row)
        if i == 0:
            set_repeat_header(row)
        for j, value in enumerate(values):
            cell = row.cells[j]
            cell.width = Cm(7.0 if j == 0 else 2.0)
            set_cell_margins(cell, top=70, start=70, bottom=70, end=70)
            if i == 0:
                set_cell_shading(cell, "EDEDED")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            r.bold = i == 0
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def answer_box(doc: Document, height_cm: float = 3.0, prompt: str | None = None) -> None:
    if prompt:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(prompt)
        r.bold = True
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="A6A6A6")
    row = table.rows[0]
    prevent_row_split(row)
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = row.cells[0]
    cell.width = Cm(17.0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    cell.paragraphs[0].add_run(" ")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_criterion(doc: Document, number: int, title: str, prompt: str) -> None:
    heading(doc, title, 1)
    p = doc.add_paragraph(prompt)
    p.paragraph_format.space_after = Pt(5)
    rating_table(doc)
    answer_box(doc, 3.6, "Komentář")
    if number in {2, 4, 6}:
        doc.add_page_break()


def checkbox_list(doc: Document, options: list[str]) -> None:
    table = doc.add_table(rows=len(options), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table, value="nil", size="0", color="FFFFFF")
    for row, option in zip(table.rows, options):
        prevent_row_split(row)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.add_run(f"[ ] {option}")


def numbered_boxes(doc: Document, count: int, height: float) -> None:
    for number in range(1, count + 1):
        answer_box(doc, height, f"{number}.")


def new_numbered_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_a4(section)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True


def build() -> Path:
    doc = Document()
    setup(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Formulář nezávislého posudku")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Experimentální rukopis pro zaslepené odborné hodnocení")
    r.italic = True
    r.font.size = Pt(11)

    field_table(doc)
    doc.add_paragraph()
    add_callout(
        doc,
        "Posudek, doporučení a známku uzavřete před vyplněním oddílu 10 o odhadu způsobu vzniku. "
        "Rukopis hodnoťte podle odborné kvality; automatický AI detektor nepoužívejte jako podklad známky.",
    )

    for idx, (title_text, prompt) in enumerate(CRITERIA, start=1):
        add_criterion(doc, idx, title_text, prompt)

    doc.add_page_break()
    heading(doc, "9. Celkový akademický závěr", 1)
    numbered_boxes(doc, 3, 1.4)
    p = doc.paragraphs[-1]
    # The preceding boxes already carry their labels; add group heading before them on next rebuild is unnecessary.

    # Insert group labels with dedicated boxes for readability.
    doc.add_paragraph().add_run("Nejsilnější stránky - použijte boxy 1-3 výše.").bold = True
    answer_box(doc, 3.0, "Nejslabší stránky")
    answer_box(doc, 3.0, "Nejzávažnější podmínka opravy před případnou obhajobou")

    heading(doc, "Doporučení", 2)
    checkbox_list(
        doc,
        [
            "doporučuji k obhajobě bez zásadní podmínky",
            "doporučuji k obhajobě po dílčích opravách",
            "doporučuji k obhajobě pouze po zásadním přepracování",
            "nedoporučuji k obhajobě",
        ],
    )

    heading(doc, "Navrhovaná známka", 2)
    checkbox_list(
        doc,
        [
            "A - výborně",
            "B - velmi dobře",
            "C - dobře",
            "D - uspokojivě",
            "E - dostatečně",
            "F - nedostatečně",
        ],
    )

    doc.add_page_break()
    heading(doc, "Otázky k obhajobě", 1)
    numbered_boxes(doc, 3, 2.3)

    heading(doc, "Dodatečné poznámky hodnotitele", 1)
    answer_box(doc, 5.0)

    doc.add_page_break()
    heading(doc, "10. Odhad způsobu vzniku", 1)
    add_callout(doc, "Tento oddíl vyplňte až po uzavření odborného posudku, doporučení a známky.")
    heading(doc, "Která možnost je podle Vás nejpravděpodobnější?", 2)
    checkbox_list(
        doc,
        [
            "převážně samostatná lidská práce bez významné generativní podpory",
            "lidská práce s běžnou jazykovou a rešeršní podporou AI",
            "lidsky řízený výzkum s významným podílem generativní AI na textu a analýze",
            "převážně generováno AI s následnou lidskou editací a kontrolou",
            "téměř plně vytvořeno generativní AI",
            "nelze rozumně odhadnout",
        ],
    )
    answer_box(doc, 1.2, "Jistota odhadu (0-100 %)")
    answer_box(doc, 4.0, "Konkrétní indicie, o které odhad opíráte")
    answer_box(doc, 3.5, "Které části působily nejvíce a nejméně autenticky?")
    answer_box(doc, 1.4, "Použil/a jste před uzavřením odhadu automatický AI detektor? ano / ne")
    answer_box(doc, 3.0, "Pokud ano, uveďte název, datum, výsledek a zda ovlivnil známku")

    doc.add_page_break()
    heading(doc, "11. Reflexe po odtajnění", 1)
    add_callout(doc, "Vyplňte samostatně až po převzetí produkčního auditu a odtajňovací zprávy.")
    answer_box(doc, 4.2, "Změnil se po odtajnění Váš názor na akademickou kvalitu rukopisu? Proč?")
    answer_box(doc, 4.2, "Které závěry posudku zůstávají platné bez ohledu na původ textu?")
    answer_box(doc, 4.2, "Co tento případ podle Vás znamená pro vedení a hodnocení kvalifikačních prací?")
    answer_box(doc, 2.4, "Souhlasíte s anonymizovaným využitím posudku ve vyhodnocení experimentu? ano / ne / s podmínkou")

    doc.save(OUTPUT)
    print(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    build()

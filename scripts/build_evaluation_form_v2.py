#!/usr/bin/env python3
"""Build and polish the structured reviewer evaluation form."""

from __future__ import annotations

from docx import Document

import build_evaluation_form as base

base.DOC_ID = "ESG-DP-2026-REVIEW"


def replace_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = True
    else:
        run = paragraph.add_run(text)
        run.bold = True


def main() -> int:
    path = base.build()
    doc = Document(path)

    section_heading = None
    strengths_label = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "9. Celkový akademický závěr":
            section_heading = paragraph
        elif text == "Nejsilnější stránky - použijte boxy 1-3 výše.":
            strengths_label = paragraph
        elif text == "Nejslabší stránky":
            replace_paragraph_text(paragraph, "Nejslabší stránky - uveďte tři hlavní")

    if section_heading is None or strengths_label is None:
        raise RuntimeError("Could not locate the overall-evaluation labels in the generated form")

    replace_paragraph_text(strengths_label, "Nejsilnější stránky - uveďte tři hlavní")
    strengths_label._p.getparent().remove(strengths_label._p)
    section_heading._p.addnext(strengths_label._p)

    for table in doc.tables:
        for row in table.rows:
            base.prevent_row_split(row)

    doc.save(path)
    print(f"Finalized: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

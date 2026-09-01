#!/usr/bin/env python3
"""Build the blinded thesis and reviewer documents from the versioned Markdown sources."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
BUILD = ROOT / "build"
BUILD.mkdir(exist_ok=True)

MAIN_FILES = [
    ROOT / "chapters/00-introduction.md",
    ROOT / "chapters/01-theoretical-framework.md",
    ROOT / "chapters/02-regulatory-context.md",
    ROOT / "chapters/03-methodology.md",
    ROOT / "chapters/04-results.md",
    ROOT / "chapters/05-discussion.md",
    ROOT / "chapters/06-conclusion.md",
    ROOT / "chapters/references.md",
]

TITLE = "Veřejně vykazovaná implementace ESG ve vybraných velkých podnicích působících v České republice"
TITLE_EN = "Publicly Reported ESG Implementation in Selected Large Enterprises Operating in the Czech Republic"
DOC_ID = "ESG-DP-2026-BLIND-01"

INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|(?<!\*)\*[^*\n]+\*(?!\*)|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Obsah bude aktualizován při sestavení PDF."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def set_repeat_together(paragraph, keep_next: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    widow = OxmlElement("w:widowControl")
    p_pr.append(widow)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.widow_control = True

    for name, size, before, after in (
        ("Title", 18, 0, 12),
        ("Subtitle", 13, 0, 8),
        ("Heading 1", 15, 18, 9),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
        if name.startswith("Heading"):
            style.paragraph_format.first_line_indent = Cm(0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "Block Quote" not in styles:
        quote = styles.add_style("Block Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["Block Quote"]
    quote.font.name = "Times New Roman"
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.paragraph_format.left_indent = Cm(1.0)
    quote.paragraph_format.right_indent = Cm(0.5)
    quote.paragraph_format.first_line_indent = Cm(0)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(6)
    quote.paragraph_format.line_spacing = 1.15

    if "Reference" not in styles:
        ref = styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        ref = styles["Reference"]
    ref.font.name = "Times New Roman"
    ref.font.size = Pt(11)
    ref.paragraph_format.left_indent = Cm(0.75)
    ref.paragraph_format.first_line_indent = Cm(-0.75)
    ref.paragraph_format.line_spacing = 1.0
    ref.paragraph_format.space_after = Pt(5)
    ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if "Caption Thesis" not in styles:
        cap = styles.add_style("Caption Thesis", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Caption Thesis"]
    cap.font.name = "Times New Roman"
    cap.font.size = Pt(10)
    cap.font.bold = True
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Experimentální rukopis pro zaslepené odborné hodnocení"
    doc.core_properties.author = "Anonymizováno pro účely hodnocení"
    doc.core_properties.last_modified_by = "Anonymizováno"
    doc.core_properties.keywords = "ESG, CSRD, ESRS, dokumentová analýza"

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([r_pr, text_el])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Liberation Mono"
            run.font.size = Pt(10)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def clean_markdown_text(text: str) -> str:
    text = text.replace("&amp;", "&")
    text = re.sub(r"<br\s*/?>", "", text, flags=re.IGNORECASE)
    return text.strip()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(width):
            value = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, clean_markdown_text(value))
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5 if width >= 5 else 9.5)
                if i == 0:
                    run.bold = True
            if i == 0:
                set_cell_shading(cell, "E7E6E6")
        if i == 0:
            repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, markdown: str, *, references: bool = False) -> None:
    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = clean_markdown_text(" ".join(line.strip() for line in paragraph_lines))
        paragraph_lines = []
        if not text:
            return
        style = "Reference" if references and not text.startswith(("##", "###")) else None
        if text.startswith("**Tabulka ") and ":" in text:
            style = "Caption Thesis"
        paragraph = doc.add_paragraph(style=style)
        if style != "Reference" and style != "Caption Thesis":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        add_inline(paragraph, text)
        set_repeat_together(paragraph)

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)), 3)
            text = clean_markdown_text(heading.group(2))
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            add_inline(paragraph, text)
            set_repeat_together(paragraph, keep_next=True)
            i += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            p = doc.add_paragraph(style="Block Quote")
            add_inline(p, clean_markdown_text(" ".join(quote_lines)))
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            text = bullet.group(1) if bullet else numbered.group(2)
            style = "List Bullet" if bullet else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, clean_markdown_text(text))
            set_repeat_together(p)
            i += 1
            continue

        paragraph_lines.append(line)
        i += 1

    flush_paragraph()


def extract_front_section(text: str, heading: str, next_heading: str | None) -> str:
    start_marker = f"# {heading}"
    start = text.index(start_marker) + len(start_marker)
    if next_heading:
        end = text.index(f"# {next_heading}", start)
    else:
        end = len(text)
    return text[start:end].strip().strip("-").strip()


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("UNIVERZITA JANA EVANGELISTY PURKYNĚ V ÚSTÍ NAD LABEM")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FAKULTA SOCIÁLNĚ EKONOMICKÁ")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE_EN)
    r.italic = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Diplomová práce - experimentální rukopis pro zaslepené odborné hodnocení")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)

    doc.add_paragraph()
    for line in (
        "Autor: anonymizováno pro účely hodnocení",
        "Studijní program: Ekonomika a management veřejného sektoru",
        "Rok: 2026",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run(line)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run("Ústí nad Labem 2026")


def add_front_matter(doc: Document) -> None:
    front = (ROOT / "chapters/front-matter.md").read_text(encoding="utf-8")
    abstract_cs = extract_front_section(front, "Abstrakt", "Abstract")
    abstract_en = extract_front_section(front, "Abstract", "Seznam zkratek")
    abbreviations = extract_front_section(front, "Seznam zkratek", None)

    add_title_page(doc)
    doc.add_page_break()

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Poznámka k rukopisu")
    note = (
        "Tento dokument je výzkumný artefakt určený výhradně k nezávislému odbornému hodnocení kvality rukopisu. "
        "Není podáván jako kvalifikační práce k získání akademického titulu. Identita autora a informace o procesu "
        "vzniku jsou v hodnoticí verzi odděleny, aby neovlivnily posouzení textu. Po uzavření hodnocení je hodnotitelům "
        "zpřístupněn úplný metodický a produkční audit."
    )
    p = doc.add_paragraph()
    add_inline(p, note)
    p = doc.add_paragraph()
    add_inline(
        p,
        "Rukopis neobsahuje čestné prohlášení studenta, podpis ani tvrzení o absolvování skutečného studijního procesu.",
    )

    doc.add_page_break()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Abstrakt")
    parse_markdown(doc, abstract_cs)

    doc.add_page_break()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Abstract")
    parse_markdown(doc, abstract_en)

    doc.add_page_break()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Seznam zkratek")
    parse_markdown(doc, abbreviations)


def start_numbered_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.0)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_page_number_start(section, 1)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.first_line_indent = Cm(0)
    run = header.add_run(f"{DOC_ID}  |  experimentální rukopis")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(89, 89, 89)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.first_line_indent = Cm(0)
    add_page_number(footer)


def build_thesis() -> Path:
    doc = Document()
    setup_document(doc)
    add_front_matter(doc)
    start_numbered_section(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Obsah")
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    toc.paragraph_format.line_spacing = 1.0
    add_toc(toc)

    for file in MAIN_FILES:
        doc.add_page_break()
        text = file.read_text(encoding="utf-8")
        parse_markdown(doc, text, references=file.name == "references.md")

    path = BUILD / f"{DOC_ID}.docx"
    doc.save(path)
    return path


def build_markdown_document(source: Path, output_name: str, title: str) -> Path:
    doc = Document()
    setup_document(doc)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    text = source.read_text(encoding="utf-8")
    # Remove the first Markdown H1 when it duplicates the document title.
    text = re.sub(r"^#\s+[^\n]+\n", "", text, count=1)
    parse_markdown(doc, text)
    doc.core_properties.title = title
    doc.core_properties.author = "Anonymizováno pro účely experimentu"
    path = BUILD / output_name
    doc.save(path)
    return path


def main() -> int:
    outputs = [
        build_thesis(),
        build_markdown_document(
            ROOT / "reviewer-packet/README.md",
            "ESG-DP-2026-reviewer-instructions.docx",
            "Pokyny pro nezávislé zaslepené hodnocení",
        ),
        build_markdown_document(
            ROOT / "reviewer-packet/evaluation-form.md",
            "ESG-DP-2026-evaluation-form.docx",
            "Formulář nezávislého posudku",
        ),
        build_markdown_document(
            ROOT / "reviewer-packet/reveal-note.md",
            "ESG-DP-2026-post-review-reveal.docx",
            "Odtajnění produkčního procesu",
        ),
    ]
    for output in outputs:
        print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
